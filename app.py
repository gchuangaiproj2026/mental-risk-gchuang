# -*- coding: utf-8 -*-
"""
国创项目：多角色高校心理健康动态监测与智能预警平台
【完善版】上传表格驱动 + 演示数据开关
- 首页 / 学生 / 教师 / 管理员 所有图表均从用户上传的 Excel/CSV 实时计算
- 未上传时可一键打开"演示数据"预览效果
- 教师模块调用 psycho_screen 算法包（贝叶斯变点 + Copula）
启动：streamlit run app.py
"""
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import os
import io
import math
import json
import time
import traceback
from datetime import datetime, timedelta
import smtplib
import requests
from email.mime.text import MIMEText

from db import (
    init_db, get_conn,
    save_self_assess, load_self_assess,
    save_screen_batch, load_latest_screen_batch,
    save_alert, load_alerts, update_alert_status,
    save_intervention, load_interventions, update_intervention_status,
    get_user_by_username, create_user, get_user_college,
)
from auth import hash_password, check_password, login_user

try:
    from psycho_screen import psycho_risk_screen
    _ALG_AVAILABLE = True
except Exception:
    psycho_risk_screen = None
    _ALG_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False

# ===== 中文字体兼容 =====
import matplotlib.font_manager as fm
def _setup_cjk_font():
    font_paths = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                fm.fontManager.addfont(fp)
            except Exception:
                pass
    candidates = ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "WenQuanYi Zen Hei"]
    available = {f.name for f in fm.fontManager.ttflist}
    for name in candidates:
        if name in available:
            plt.rcParams['font.sans-serif'] = [name, 'DejaVu Sans']
            break
    plt.rcParams['axes.unicode_minus'] = False
_setup_cjk_font()

st.set_page_config(page_title="高校心理健康智能监测平台", layout="wide", initial_sidebar_state="collapsed")

# ===== 全局CSS =====
st.markdown("""
<style>
.block-container { padding-top: 1rem !important; padding-left: 1rem; padding-right: 1rem; max-width: 1450px; }
.dash-title { font-size:26px; font-weight:bold; color:#0F4C99; margin-bottom:20px; }
.sub-title { font-size: 20px; font-weight: 700; color: #1e293b; margin: .8rem 0 1rem; }
.res-card { background:#fff; border:1px solid #eef2f7; border-radius:16px; padding:1.4rem 1.5rem; box-shadow:0 3px 10px rgba(0,0,0,.05); height:100%; margin-bottom:1rem; }
.top-bar-wrap{margin-top:28px;margin-bottom:12px;}
.top-bar{background:linear-gradient(135deg,#0F4C99,#1a6bc4);color:white;padding:14px 24px;border-radius:14px;display:flex;justify-content:space-between;align-items:center;box-shadow:0 4px 14px rgba(15,76,153,0.2);}
.top-logo{font-size:22px;font-weight:bold;display:flex;align-items:center;gap:10px;}
.top-user-info{display:flex;gap:18px;align-items:center;font-size:15px;}
.logout-btn{background:rgba(255,255,255,0.2);border:1px solid rgba(255,255,255,0.28);color:white;padding:7px 18px;border-radius:8px;cursor:pointer;text-decoration:none;}
.logout-btn:hover{background:rgba(255,255,255,0.32);}
@media (max-width: 768px) {
    .block-container { padding-left: 0.6rem; padding-right: 0.6rem; padding-top:1.2rem !important; }
    .top-bar{flex-direction:column;gap:10px;padding:12px 16px;}
}
</style>
""", unsafe_allow_html=True)

# ===== 通用卡片 =====
def _metric_card_html(title, value, delta="", color="#0F4C99", icon=""):
    return (
        '<div style="background:{c};border-radius:14px;padding:20px 24px;color:white;'
        'box-shadow:0 6px 16px rgba(0,0,0,0.16);height:100%;">'
        '<div style="font-size:15px;opacity:0.88;display:flex;justify-content:space-between;">'
        '<span>{t}</span><span style="font-size:28px;">{i}</span></div>'
        '<div style="font-size:34px;font-weight:bold;margin:8px 0;">{v}</div>'
        '<div style="font-size:14px;opacity:0.92;">{d}</div></div>'
    ).format(c=color, t=title, i=icon, v=value, d=delta)

def _suggestion_card_html(title, desc, action, color="#4f6ef7"):
    return (
        '<div style="background:{c};border-radius:14px;padding:20px 24px;color:white;'
        'margin-bottom:16px;box-shadow:0 4px 12px rgba(0,0,0,0.12);">'
        '<div style="font-weight:bold;font-size:18px;display:flex;align-items:center;gap:10px;">{t}</div>'
        '<div style="font-size:16px;opacity:0.95;margin:6px 0;">{d}</div>'
        '<div style="font-size:15px;opacity:0.85;margin-top:6px;">💡 {a}</div></div>'
    ).format(c=color, t=title, d=desc, a=action)

# ===== 消息推送 =====
def send_email(recipient, subject, body):
    try:
        sender = os.getenv("SMTP_USER", "alert@school.edu")
        password = os.getenv("SMTP_PASS", "")
        if not password:
            st.warning("未配置邮件密码，消息未发送")
            return
        msg = MIMEText(body, "plain", "utf-8")
        msg["Subject"] = subject
        msg["From"] = sender
        msg["To"] = recipient
        server = smtplib.SMTP(os.getenv("SMTP_HOST", "smtp.qq.com"), int(os.getenv("SMTP_PORT", 587)))
        server.starttls()
        server.login(sender, password)
        server.sendmail(sender, [recipient], msg.as_string())
        server.quit()
        st.success("邮件已发送")
    except Exception as e:
        st.error(f"邮件发送失败: {e}")

def send_wechat_robot(webhook_url, content):
    try:
        data = {"msgtype": "text", "text": {"content": content}}
        response = requests.post(webhook_url, json=data)
        if response.status_code == 200:
            st.success("企业微信通知已发送")
        else:
            st.error(f"通知失败: {response.text}")
    except Exception as e:
        st.error(f"通知异常: {e}")

# ===================== 数据层：上传解析 =====================
RISK_COLORS = {"高危": "#D93025", "中危": "#F57C00", "低危": "#2E7D32",
               "低风险": "#2E7D32", "预警": "#D93025", "关注": "#F57C00", "良好": "#2E7D32"}

def parse_upload(uploaded_file):
    """把用户上传的 Excel/CSV 解析为 DataFrame（自动处理编码）"""
    name = uploaded_file.name.lower()
    try:
        if name.endswith(".xlsx") or name.endswith(".xls"):
            return pd.read_excel(uploaded_file)
        if name.endswith(".csv"):
            raw = uploaded_file.getvalue()
            for enc in ("utf-8-sig", "gb18030", "gbk", "utf-8"):
                try:
                    return pd.read_csv(io.BytesIO(raw), encoding=enc, sep=None, engine="python")
                except UnicodeDecodeError:
                    continue
            return pd.read_csv(io.BytesIO(raw), encoding="utf-8", errors="replace")
    except Exception as e:
        st.error(f"文件解析失败（{name}）：{e}")
        return None
    st.error("仅支持 Excel(.xlsx/.xls) 或 CSV 文件")
    return None

def classify_from_score(s, high_thr=1.5, mid_thr=0.5):
    """根据综合风险分与阈值划分等级（不重新采样）"""
    return np.where(s >= high_thr, "高危", np.where(s >= mid_thr, "中危", "低危"))

def quick_risk(df):
    """快速分级：无需跑 MCMC，仅用于首页/管理员的实时概览"""
    df = df.copy()
    if "风险等级" in df.columns:
        return df
    score_col = next((c for c in ["综合风险分", "总分", "score"] if c in df.columns), None)
    if score_col is None:
        num = df.select_dtypes(include=np.number)
        if num.shape[1] == 0:
            return df
        df["总分"] = num.sum(axis=1)
        score_col = "总分"
    z = (df[score_col] - df[score_col].mean()) / (df[score_col].std(ddof=0) or 1.0)
    df["综合风险分"] = z
    df["风险等级"] = classify_from_score(
        z, st.session_state.get("global_high_thr", 1.5), st.session_state.get("global_mid_thr", 0.5))
    return df

KNOWN_DIMS = ["焦虑", "抑郁", "压力", "睡眠障碍", "社交回避", "学业压力", "人际关系", "情绪波动", "生活满意度"]

def normalize_student_cols(df):
    """学生自评表列名兼容（中文/英文）"""
    mapping = {
        "time": "自评时间", "date": "自评时间", "日期": "自评时间",
        "anxiety": "焦虑", "depression": "抑郁", "stress": "压力",
        "sleep": "睡眠障碍", "睡眠": "睡眠障碍", "social": "社交回避", "社交": "社交回避",
    }
    df = df.rename(columns={k: v for k, v in mapping.items() if k in df.columns})
    need = ["焦虑", "抑郁", "压力", "睡眠障碍", "社交回避"]
    if "总分" not in df.columns and all(c in df.columns for c in need):
        df["总分"] = df[need].sum(axis=1)
    return df

# ===================== 演示数据生成 =====================
def demo_toggle_widget(help_text="未上传数据时可打开，预览各角色页面效果"):
    st.toggle("🧪 演示数据预览", key="use_demo", help=help_text)

def demo_survey(n=150, seed=42):
    rng = np.random.default_rng(seed)
    names = ["张伟", "李娜", "王芳", "刘洋", "陈静", "杨帆", "赵磊", "孙悦", "周婷", "吴昊", "郑爽", "钱进", "冯雪", "韩磊", "曹阳"]
    schools = ["计算机学院", "外国语学院", "经济管理学院", "机械工程学院", "文学院", "艺术学院"]
    grades = ["大一", "大二", "大三", "大四"]
    df = pd.DataFrame({
        "学号": [f"2026{1000+i}" for i in range(n)],
        "姓名": [rng.choice(names) for _ in range(n)],
        "学院": [rng.choice(schools) for _ in range(n)],
        "年级": [rng.choice(grades, p=[0.3, 0.3, 0.25, 0.15]) for _ in range(n)],
        "焦虑": np.clip(rng.normal(13, 6, n), 2, 38).astype(int),
        "抑郁": np.clip(rng.normal(11, 5.5, n), 2, 36).astype(int),
        "压力": np.clip(rng.normal(15, 6, n), 3, 40).astype(int),
        "睡眠障碍": np.clip(rng.normal(8, 5, n), 0, 32).astype(int),
        "社交回避": np.clip(rng.normal(7, 4.5, n), 0, 30).astype(int),
    })
    for _ in range(int(n * 0.06)):
        idx = rng.integers(0, n)
        df.loc[idx, ["焦虑", "抑郁", "压力", "睡眠障碍", "社交回避"]] += rng.integers(15, 30, size=5)
    return df

def demo_student_history(seed=7):
    rng = np.random.default_rng(seed)
    dates = pd.date_range(end=datetime.now(), periods=8, freq="14D")
    rows = []
    base = [8, 9, 10, 6, 5]
    for i, d in enumerate(dates):
        vals = [max(0, int(base[j] + rng.normal(0, 3) + (2 if i >= 4 else 0))) for j in range(5)]
        rows.append({"自评时间": d.strftime("%Y-%m-%d %H:%M"),
                     "焦虑": vals[0], "抑郁": vals[1], "压力": vals[2],
                     "睡眠障碍": vals[3], "社交回避": vals[4]})
    df = pd.DataFrame(rows)
    df["总分"] = df[["焦虑", "抑郁", "压力", "睡眠障碍", "社交回避"]].sum(axis=1)
    return df

def demo_alert_df(seed=9):
    rng = np.random.default_rng(seed)
    names = ["张伟", "李娜", "王芳", "刘洋", "陈静", "杨帆", "赵磊", "孙悦", "周婷", "吴昊", "郑爽", "钱进"]
    schools = ["计算机学院", "外国语学院", "经济管理学院", "机械工程学院"]
    n = 12
    return pd.DataFrame({
        "学号": [f"2026{1000+i}" for i in range(1, n+1)],
        "姓名": [rng.choice(names) for _ in range(n)],
        "学院": [rng.choice(schools) for _ in range(n)],
        "风险等级": [rng.choice(["高危", "高危", "中危"]) for _ in range(n)],
        "综合风险分": [round(float(rng.uniform(1.5, 3.2)), 2) for _ in range(n)],
        "预警时间": [(datetime.now() - timedelta(days=int(rng.integers(0, 12)))).strftime("%Y-%m-%d") for _ in range(n)],
        "处置状态": [rng.choice(["待跟进", "干预中", "已转介", "已结案"], p=[0.3, 0.4, 0.2, 0.1]) for _ in range(n)],
    })

def demo_intervention_df(seed=21):
    rng = np.random.default_rng(seed)
    names = ["张伟", "李娜", "王芳", "刘洋", "陈静", "杨帆", "赵磊", "孙悦"]
    return pd.DataFrame({
        "学号": [f"2026{1000+i}" for i in range(1, 9)],
        "学生": [rng.choice(names) for _ in range(8)],
        "学院": [rng.choice(["计算机学院", "外国语学院", "经济管理学院", "机械工程学院"]) for _ in range(8)],
        "风险等级": [rng.choice(["高危", "高危", "中危"]) for _ in range(8)],
        "干预方式": [rng.choice(["一对一咨询", "正念减压", "团体辅导", "家校联动", "转介中心"]) for _ in range(8)],
        "负责人": [rng.choice(["张老师", "李老师", "王老师", "刘老师"]) for _ in range(8)],
        "开始时间": [(datetime.now() - timedelta(days=int(rng.integers(1, 30)))).strftime("%Y-%m-%d") for _ in range(8)],
        "状态": [rng.choice(["干预中", "干预中", "已结案"], p=[0.5, 0.3, 0.2]) for _ in range(8)],
    })

def demo_roster():
    df = demo_survey(n=120, seed=3)
    df = quick_risk(df)
    df["心理状态"] = df["风险等级"].map({"高危": "预警", "中危": "关注", "低危": "良好"}).fillna("良好")
    days = np.random.default_rng(5).integers(1, 90, len(df))
    df["最近测评"] = [(datetime.now() - timedelta(days=int(x))).strftime("%Y-%m-%d") for x in days]
    return df[["学号", "姓名", "学院", "年级", "心理状态", "最近测评", "综合风险分"]]

# ===================== 图表函数（全部由 DataFrame 计算） =====================
def chart_risk_pie(df, col="风险等级"):
    counts = df[col].value_counts()
    order = [c for c in ["高危", "中危", "低危", "低风险", "预警", "关注", "良好"] if c in counts.index]
    labels = order or list(counts.index)
    vals = [int(counts[c]) for c in labels]
    colors = [RISK_COLORS.get(c, "#94a3b8") for c in labels]
    fig, ax = plt.subplots(figsize=(9, 5), dpi=110)
    wedges, _, autotexts = ax.pie(vals, labels=labels, colors=colors, autopct="%.1f%%",
                                  startangle=90, textprops={"fontsize": 11})
    for w in wedges:
        w.set_edgecolor("white"); w.set_linewidth(1.5)
    ax.set_title(f"{col}分布（总人数 {int(counts.sum()):,}）", fontsize=12)
    fig.tight_layout()
    return fig

def chart_risk_bar(df, col="风险等级"):
    counts = df[col].value_counts()
    order = [c for c in ["高危", "中危", "低危", "低风险", "预警", "关注", "良好"] if c in counts.index]
    labels = order or list(counts.index)
    vals = [int(counts[c]) for c in labels]
    colors = [RISK_COLORS.get(c, "#94a3b8") for c in labels]
    fig, ax = plt.subplots(figsize=(9, 5), dpi=110)
    bars = ax.bar(labels, vals, color=colors, width=0.5)
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width()/2, b.get_height() + max(vals)*0.01,
                str(v), ha="center", fontsize=11, fontweight="bold")
    ax.set_ylabel("人数")
    ax.tick_params(axis="both", labelsize=11)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

def chart_hist(df, col="总分"):
    fig, ax = plt.subplots(figsize=(9, 5), dpi=110)
    ax.hist(df[col], bins=20, color="#1a6bc4", edgecolor="white", alpha=0.85)
    ax.axvline(df[col].mean(), color="#D93025", linestyle="--", linewidth=1.8,
               label=f"均值 {df[col].mean():.1f}")
    ax.set_xlabel(col); ax.set_ylabel("人数"); ax.legend()
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

def chart_radar(labels, values, title="", color="#0F4C99"):
    N = len(labels)
    angles = [n / float(N) * 2 * math.pi for n in range(N)]
    vals = values + [values[0]]
    angs = angles + [angles[0]]
    fig = plt.figure(figsize=(7, 5.5), dpi=110)
    ax = fig.add_subplot(111, polar=True)
    ax.plot(angs, vals, color=color, linewidth=2)
    ax.fill(angs, vals, color=color, alpha=0.2)
    ax.set_xticks(angles)
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_yticks([20, 40, 60, 80])
    ax.set_ylim(0, 100)
    ax.set_title(title, fontsize=12)
    return fig

def chart_multi_trend(df, dims, time_col="自评时间"):
    t = pd.to_datetime(df[time_col], errors="coerce")
    fig, ax = plt.subplots(figsize=(10, 5), dpi=110)
    for d in dims:
        ax.plot(t, df[d], marker="o", markersize=4, linewidth=1.8, label=d)
    ax.legend(ncol=3, fontsize=10)
    ax.set_ylabel("得分")
    ax.tick_params(axis="x", rotation=30, labelsize=9)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

def chart_copula_scatter(df, x_col="Copula异常分", y_col="综合风险分", level_col="风险等级"):
    fig, ax = plt.subplots(figsize=(10, 5.5), dpi=110)
    if x_col in df.columns and y_col in df.columns:
        for level, color in RISK_COLORS.items():
            sub = df[df[level_col] == level]
            if len(sub):
                ax.scatter(sub[x_col], sub[y_col], s=26, c=color, label=level, alpha=0.75)
        ax.set_xlabel(x_col); ax.set_ylabel(y_col)
        ax.legend(fontsize=10)
    else:
        num = df.select_dtypes(include=np.number)
        if num.shape[1] < 2:
            ax.text(0.5, 0.5, "数值列不足", ha="center", va="center")
        else:
            xc, yc = num.columns[0], num.columns[1]
            sc = ax.scatter(df[xc], df[yc], c=df[level_col].map(RISK_COLORS), s=26, alpha=0.75)
            ax.set_xlabel(xc); ax.set_ylabel(yc)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

def chart_dims_box(df, dims):
    fig, ax = plt.subplots(figsize=(10, 5), dpi=110)
    df[dims].boxplot(ax=ax, patch_artist=True,
                     medianprops=dict(color="#D93025", linewidth=1.6))
    ax.set_ylabel("得分")
    ax.tick_params(axis="x", rotation=20)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

def chart_corr(df, dims):
    fig, ax = plt.subplots(figsize=(7, 6), dpi=110)
    corr = df[dims].corr()
    im = ax.imshow(corr, cmap="RdBu_r", vmin=-1, vmax=1)
    ax.set_xticks(range(len(dims))); ax.set_xticklabels(dims, rotation=45, ha="right")
    ax.set_yticks(range(len(dims))); ax.set_yticklabels(dims)
    for i in range(len(dims)):
        for j in range(len(dims)):
            ax.text(j, i, f"{corr.iloc[i, j]:.2f}", ha="center", va="center", fontsize=9)
    fig.colorbar(im, ax=ax)
    fig.tight_layout()
    return fig

def chart_tau_series(df, tau, score_col="总分"):
    fig, ax = plt.subplots(figsize=(11, 4.5), dpi=110)
    x = np.arange(len(df))
    ax.plot(x, df[score_col], color="#1a6bc4", linewidth=1.2, alpha=0.85)
    if tau is not None:
        ax.axvline(tau, color="#D93025", linestyle="--", linewidth=2,
                   label=f"检测变点 τ≈{tau:.0f}")
    ax.set_xlabel("样本序号"); ax.set_ylabel(score_col)
    ax.legend(fontsize=10)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

# ===================== Session 初始化 =====================
def init_session():
    init_db()
    defaults = {
        "page_root": "home", "role": "", "username": "", "user_college": None,
        "use_demo": False,
        "home_df": None,
        "stu_history_df": None,
        "screen_raw_df": None, "screen_result": None, "screen_tau": None,
        "screen_copula": None, "screen_dims": [], "screen_batch_name": "",
        "admin_roster_df": None, "admin_alert_df": None, "admin_inter_df": None,
        "global_high_thr": 1.5, "global_mid_thr": 0.5,
    }
    for k, v in defaults.items():
        st.session_state.setdefault(k, v)

def load_user_data(username, role, college):
    """登录后加载该用户的历史数据（学生自评 / 教师最新批次）"""
    if role == "student":
        hist = load_self_assess(username, college=college)
        if hist is not None and not hist.empty:
            st.session_state["stu_history_df"] = hist
    elif role in ("teacher", "admin"):
        batch_name, tau, df_data, copula = load_latest_screen_batch(college=college if role == "teacher" else None)
        if df_data is not None:
            st.session_state["screen_result"] = df_data
            st.session_state["screen_tau"] = tau
            st.session_state["screen_copula"] = copula
            st.session_state["screen_batch_name"] = batch_name

# ===================== 顶部导航 =====================
def render_top_nav_home():
    qp = st.query_params
    current_tab = qp.get("tab", "overview")
    tabs = [
        ("overview", "📊 总览大屏"),
        ("survey", "📝 心理普查"),
        ("analysis", "📈 数据分析"),
        ("resources", "📚 资源中心"),
    ]
    nav_items = ""
    for key, label in tabs:
        active_cls = "active" if current_tab == key else ""
        nav_items += f'<a href="?tab={key}" class="nav-item-home {active_cls}">{label}</a>'
    nav_html = f"""
    <style>
        .top-nav-home {{ background: linear-gradient(135deg, #0F4C99, #1a6bc4); padding: 0 3rem; display: flex;
            align-items: center; justify-content: space-between; height: 84px;
            box-shadow: 0 6px 20px rgba(0,0,0,0.20); border-radius: 16px; margin-top: 40px; margin-bottom:12px;
            border: 1px solid rgba(255,255,255,0.10); }}
        .nav-brand {{ color: white; font-size: 26px; font-weight: 700; display: flex; align-items: center; gap: 14px; white-space: nowrap; }}
        .nav-brand .sub {{ font-size: 15px; font-weight: 400; opacity: 0.75; margin-left: 6px; }}
        .nav-menu-home {{ display: flex; gap:10px; flex: 1; margin: 0 2.5rem; }}
        .nav-item-home {{ color: #ffffff; padding: 12px 22px; border-radius:10px; font-size:17px;
            white-space: nowrap; text-decoration: none; transition: all 0.3s; font-weight: 500; cursor:pointer; }}
        .nav-item-home:hover {{ color: #ffffff; background: rgba(255,255,255,0.25); }}
        .nav-item-home.active {{ color: #0F4C99; background: #ffffff; font-weight: 700; box-shadow: 0 2px 8px rgba(0,0,0,0.15); }}
        .nav-user-home {{ color: rgba(255,255,255,0.92); font-size:16px; display: flex; align-items: center; gap:22px; }}
        .login-btn {{ background: rgba(255,255,255,0.22); border: 1px solid rgba(255,255,255,0.30); color: white;
            padding:10px 28px; border-radius:10px; cursor: pointer; font-size:16px; transition: all 0.3s;
            text-decoration: none; font-weight: 500; }}
        .login-btn:hover {{ background: rgba(255,255,255,0.36); transform: translateY(-2px); box-shadow:0 6px 16px rgba(0,0,0,0.22); }}
        .nav-divider {{ color: rgba(255,255,255,0.25); font-size:24px; }}
        @media (max-width: 768px) {{
            .top-nav-home {{ padding: 0 1rem; flex-wrap: wrap; height: auto; min-height:70px; margin-top:20px;
                border-radius:10px; padding:0.8rem 1rem; }}
            .nav-brand {{ font-size:20px; }} .nav-brand .sub {{ display:none; }}
            .nav-item-home {{ font-size:14px; padding:8px 14px; }}
            .nav-user-home {{ font-size:14px; gap:12px; }} .login-btn {{ padding:7px 16px; font-size:14px; }}
            .nav-menu-home {{ margin:0 0.4rem; gap:6px; flex-wrap:wrap; }} .nav-divider {{ display:none; }}
        }}
    </style>
    <div class="top-nav-home">
        <div class="nav-brand">🧠 心理监测平台 <span class="sub">| 健康校园 · 智能预警</span></div>
        <div class="nav-menu-home">
            {nav_items}
        </div>
        <div class="nav-user-home">
            <span style="opacity:0.65;">👤 访客</span>
            <span class="nav-divider">|</span>
            <a href="?login=1" class="login-btn">🔑 登录</a>
        </div>
    </div>
    """
    st.markdown(nav_html, unsafe_allow_html=True)
    if "login" in qp:
        st.session_state["page_root"] = "role_login"
        st.query_params.clear()
        st.rerun()

def render_top_bar():
    username = st.session_state.get("username", "")
    role = st.session_state.get("role", "")
    college = st.session_state.get("user_college", "")
    html = f"""
    <div class="top-bar-wrap">
        <div class="top-bar">
            <div class="top-logo">🧠 高校心理健康动态监测与智能预警平台</div>
            <div class="top-user-info">
                <span>👤 {username}｜{role}｜🏫{college if college else '全校'}</span>
                <a href="?logout=1" class="logout-btn">🚪退出登录</a>
            </div>
        </div>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)
    qp = st.query_params
    if "logout" in qp:
        st.session_state["page_root"] = "home"
        st.session_state["role"] = ""
        st.session_state["username"] = ""
        st.session_state["user_college"] = None
        st.query_params.clear()
        st.rerun()

# ===================== 访客首页 =====================
def render_home_page():
    render_top_nav_home()
    tab = st.query_params.get("tab", "overview")

    # ===== 心理普查页 =====
    if tab == "survey":
        st.markdown("""
        <div style="text-align:center;padding:0.5rem 0 1rem 0;">
            <span style="font-size:32px;font-weight:bold;color:#0F4C99;">📝 心理普查</span>
            <span style="font-size:18px;color:#666;margin-left:20px;">面向全体学生的心理健康状况普查</span>
        </div>
        """, unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown(_metric_card_html("普查量表", "SCL-90 / PHQ-9", "焦虑·抑郁·压力", "#0F4C99", "📋"), unsafe_allow_html=True)
        with c2:
            st.markdown(_metric_card_html("筛查算法", "贝叶斯+Copula", "变点检测·相依分析", "#1a6bc4", "🧠"), unsafe_allow_html=True)
        with c3:
            st.markdown(_metric_card_html("风险分级", "高危/中危/低危", "三级预警机制", "#2E7D32", "🚦"), unsafe_allow_html=True)
        st.divider()
        st.markdown("### 普查流程")
        steps = [
            ("1️⃣ 学生测评", "学生在线填写心理量表，数据实时入库"),
            ("2️⃣ 智能筛查", "教师上传普查数据，系统自动运行贝叶斯变点检测与Copula相依分析"),
            ("3️⃣ 风险分级", "按综合风险分自动划分高危、中危、低危三个等级"),
            ("4️⃣ 干预跟进", "高危学生自动进入预警台账，辅导员一对一跟进"),
        ]
        for title, desc in steps:
            st.markdown(f"- **{title}**：{desc}")
        st.divider()
        st.info("💡 教师登录后可在「教师端 → 上传与筛查」中上传普查数据表并运行智能筛查算法。")
        if st.button("🔑 登录进入教师端", use_container_width=True):
            st.session_state["page_root"] = "role_login"
            st.rerun()
        return

    # ===== 数据分析页 =====
    if tab == "analysis":
        st.markdown("""
        <div style="text-align:center;padding:0.5rem 0 1rem 0;">
            <span style="font-size:32px;font-weight:bold;color:#0F4C99;">📈 数据分析</span>
            <span style="font-size:18px;color:#666;margin-left:20px;">多维度心理健康数据可视化分析</span>
        </div>
        """, unsafe_allow_html=True)
        demo = st.session_state.get("use_demo", False)
        col_u, col_d = st.columns([3, 1])
        with col_u:
            up = st.file_uploader("📤 上传测评数据表进行分析", type=["xlsx", "xls", "csv"], key="analysis_up")
        with col_d:
            demo_toggle_widget("打开演示数据查看分析效果")
        df = None
        if up is not None:
            df = parse_upload(up)
        if df is None and demo:
            df = quick_risk(demo_survey())
            st.info("当前展示【演示数据】分析结果")
        if df is None:
            st.info("👆 请上传测评数据表，或打开演示数据预览分析图表。")
            return
        df = quick_risk(df)
        dims = [c for c in KNOWN_DIMS if c in df.columns]
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**📊 风险等级分布**")
            st.pyplot(chart_risk_pie(df))
        with c2:
            st.markdown("**📈 综合风险分分布**")
            st.pyplot(chart_hist(df, "综合风险分" if "综合风险分" in df.columns else "总分"))
        st.divider()
        if len(dims) >= 2:
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**📦 各维度得分箱线图**")
                st.pyplot(chart_dims_box(df, dims))
            with c2:
                st.markdown("**🔗 维度相关性热力图**")
                st.pyplot(chart_corr(df, dims))
        st.divider()
        st.markdown("**🔍 风险散点分布**")
        st.pyplot(chart_copula_scatter(df))
        return

    # ===== 资源中心页 =====
    if tab == "resources":
        st.markdown("""
        <div style="text-align:center;padding:0.5rem 0 1rem 0;">
            <span style="font-size:32px;font-weight:bold;color:#0F4C99;">📚 资源中心</span>
            <span style="font-size:18px;color:#666;margin-left:20px;">心理健康科普与自助资源</span>
        </div>
        """, unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("### 🧘 自助调节")
            st.markdown("""
            - **正念冥想**：每天10分钟正念练习，缓解焦虑情绪
            - **睡眠卫生**：规律作息，睡前1小时远离电子屏幕
            - **运动减压**：每周3次有氧运动，每次30分钟以上
            - **社交支持**：主动与家人朋友交流，避免独处封闭
            """)
            st.markdown("### 📞 求助渠道")
            st.markdown("""
            - **校内心理咨询中心**：预约一对一专业咨询
            - **全国心理援助热线**：400-161-9995
            - **北京心理危机研究与干预中心**：010-82951332
            - **希望24热线**：400-161-9995
            """)
        with c2:
            st.markdown("### 📖 科普知识")
            st.markdown("""
            **焦虑症常见表现**
            - 持续紧张、坐立不安
            - 心悸、出汗、手抖
            - 睡眠困难、注意力不集中

            **抑郁症常见表现**
            - 持续情绪低落、兴趣减退
            - 精力下降、疲劳感
            - 自我评价低、自责
            - 睡眠和食欲改变

            **何时需要专业帮助**
            - 症状持续2周以上
            - 严重影响学习、生活和社交
            - 出现自伤或自杀念头
            """)
        st.divider()
        st.warning("⚠️ 本平台资源仅供科普参考，不能替代专业医疗诊断。如有需要请及时就医或联系专业心理咨询师。")
        return

    # ===== 总览大屏（默认页） =====
    st.markdown("""
    <div style="text-align:center;padding:0.5rem 0 1rem 0;">
        <span style="font-size:32px;font-weight:bold;color:#0F4C99;">🏠 心理健康监测总览</span>
        <span style="font-size:18px;color:#666;margin-left:20px;">数据均来自你上传的测评表</span>
    </div>
    """, unsafe_allow_html=True)
    demo = st.session_state.get("use_demo", False)
    col_u, col_d = st.columns([3, 1])
    with col_u:
        up = st.file_uploader("📤 上传测评数据表（Excel/CSV，含各维度得分或总分/风险等级列即可）",
                              type=["xlsx", "xls", "csv"], key="home_up")
    with col_d:
        demo_toggle_widget("未上传数据时可打开，用模拟测评数据预览图表")

    df = st.session_state.get("home_df")
    if up is not None:
        df = parse_upload(up)
        if df is not None:
            st.session_state["home_df"] = df

    if df is None and demo:
        df = quick_risk(demo_survey())
        st.info("当前展示的是【演示数据】，上传你的表格后会自动切换为真实数据")

    if df is None:
        st.info("👆 请先上传测评数据表（或打开右上角「演示数据预览」），首页所有指标与图表将实时计算生成。")
        return

    df = quick_risk(df)
    dims = [c for c in KNOWN_DIMS if c in df.columns]
    has_level = "风险等级" in df.columns

    if has_level:
        high_n = int((df["风险等级"] == "高危").sum())
        mid_n = int((df["风险等级"] == "中危").sum())
        ratio = high_n / len(df) * 100 if len(df) else 0
        cards = [
            {"title": "测评总人数", "value": f"{len(df):,}", "delta": f"{df.shape[1]} 个字段", "color": "#0F4C99", "icon": "👥"},
            {"title": "高危人数", "value": f"{high_n:,}", "delta": f"占比 {ratio:.1f}%", "color": "#D93025", "icon": "🚨"},
            {"title": "中危人数", "value": f"{mid_n:,}", "delta": f"需重点关注", "color": "#F57C00", "icon": "⚠️"},
            {"title": "低危/正常", "value": f"{len(df)-high_n-mid_n:,}", "delta": "保持良好状态", "color": "#2E7D32", "icon": "✅"},
        ]
        cols = st.columns(4)
        for i, c in enumerate(cols):
            with c:
                st.markdown(_metric_card_html(**cards[i]), unsafe_allow_html=True)
        st.divider()

        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**📊 预警分层统计**")
            st.pyplot(chart_risk_pie(df))
        with c2:
            st.markdown("**📊 风险等级人数**")
            st.pyplot(chart_risk_bar(df))
        st.divider()

        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**📈 综合风险分分布**")
            st.pyplot(chart_hist(df, "综合风险分" if "综合风险分" in df.columns else "总分"))
        with c2:
            if dims:
                st.markdown("**🕸️ 各维度均值雷达图**")
                vals = [float(df[d].mean() / max(40, df[d].max()) * 100) for d in dims]
                st.pyplot(chart_radar(vals, dims, title="风险维度均值（归一化 0-100）"))
        st.divider()

        c1, c2 = st.columns([3, 2])
        with c1:
            st.markdown("**🔍 风险散点分布**")
            st.pyplot(chart_copula_scatter(df))
        with c2:
            st.markdown("**🚨 预警动态（风险分最高 5 人）**")
            warn = df.sort_values("综合风险分", ascending=False).head(5)
            cols_w = [c for c in ["学号", "姓名", "学院", "风险等级", "综合风险分"] if c in warn.columns]
            st.dataframe(warn[cols_w].style.map(
                lambda x: "color:#D93025;font-weight:bold" if x == "高危" else "",
                subset=["风险等级"] if "风险等级" in cols_w else []),
                hide_index=True, use_container_width=True)
        if len(dims) >= 2:
            with st.expander("📎 维度相关性热力图"):
                st.pyplot(chart_corr(df, dims))
    else:
        st.warning("未识别到风险等级列。请上传包含总分/维度得分的量表数据（列如：焦虑、抑郁、压力、睡眠障碍、社交回避）。")
        st.markdown(f"**已读取 {len(df)} 行 × {df.shape[1]} 列，字段：{list(df.columns)}**")
        num = df.select_dtypes(include=np.number)
        if num.shape[1] >= 2:
            c1, c2 = st.columns(2)
            with c1:
                st.pyplot(chart_hist(df.assign(总分=num.sum(axis=1)), "总分"))
            with c2:
                st.pyplot(chart_corr(df, list(num.columns)[:8]))

# ===================== 登录页 =====================
def render_login_page():
    st.markdown("""
    <div style="text-align:center;margin-top:50px;">
        <div style="font-size:40px;">🧠</div>
        <div style="font-size:28px;font-weight:bold;color:#0F4C99;margin:8px 0;">高校心理健康动态监测与智能预警平台</div>
        <div style="font-size:15px;color:#666;">多角色登录 · 上传数据 · 智能筛查</div>
    </div>
    """, unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c2:
        with st.form("login_form"):
            username = st.text_input("用户名")
            password = st.text_input("密码", type="password")
            submitted = st.form_submit_button("🔑 登录", use_container_width=True)
        if submitted:
            role, college = login_user(username, password)
            if role:
                st.session_state["role"] = role
                st.session_state["username"] = username
                st.session_state["user_college"] = college
                st.session_state["page_root"] = "app"
                load_user_data(username, role, college)
                st.rerun()
            else:
                st.error("用户名或密码错误")
        st.info("演示账号：admin / admin123（管理员）· teacher / 123456（教师）· student / 123456（学生）")
        if st.button("← 返回首页"):
            st.session_state["page_root"] = "home"
            st.rerun()

# ===================== 学生模块 =====================
def _merge_stu_history(df_new):
    cur = st.session_state.get("stu_history_df")
    if cur is None or cur.empty:
        st.session_state["stu_history_df"] = df_new.reset_index(drop=True)
    else:
        st.session_state["stu_history_df"] = pd.concat([cur, df_new], ignore_index=True)

def render_student_page():
    render_top_bar()
    st.markdown('<div class="sub-title">🧑‍🎓 学生自助心理测评与自评报告</div>', unsafe_allow_html=True)
    demo_toggle_widget("打开后无需上传即可看到模拟自评历史")
    username = st.session_state.get("username", "")
    tabs = st.tabs(["📝 在线自评", "📤 导入历史测评", "📊 我的报告"])
    with tabs[0]:
        with st.form("self_assess_form"):
            c1, c2 = st.columns(2)
            with c1:
                anxiety = st.slider("焦虑（0-40）", 0, 40, 10)
                depression = st.slider("抑郁（0-40）", 0, 40, 10)
                stress = st.slider("压力（0-40）", 0, 40, 12)
            with c2:
                sleep = st.slider("睡眠障碍（0-40）", 0, 40, 8)
                social = st.slider("社交回避（0-40）", 0, 40, 8)
            note = st.text_area("补充说明（选填）", "")
            submitted = st.form_submit_button("💾 提交自评", use_container_width=True)
        if submitted:
            row = pd.DataFrame([{
                "自评时间": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "焦虑": anxiety, "抑郁": depression, "压力": stress,
                "睡眠障碍": sleep, "社交回避": social,
            }])
            row["总分"] = row[["焦虑", "抑郁", "压力", "睡眠障碍", "社交回避"]].sum(axis=1)
            try:
                save_self_assess(username, row, st.session_state.get("user_college"))
                _merge_stu_history(row)
                st.success("✅ 自评已保存，请到「我的报告」查看趋势")
            except Exception as e:
                st.error(f"保存失败：{e}")

    with tabs[1]:
        up = st.file_uploader("📤 上传自评历史表（列名：自评时间 / 焦虑 / 抑郁 / 压力 / 睡眠障碍 / 社交回避，中英文均可）",
                              type=["xlsx", "xls", "csv"], key="stu_up")
        if up is not None:
            df = parse_upload(up)
            if df is not None:
                df = normalize_student_cols(df)
                need = ["自评时间", "焦虑", "抑郁", "压力", "睡眠障碍", "社交回避"]
                missing = [c for c in need if c not in df.columns]
                if missing:
                    st.error(f"缺少列：{missing}")
                else:
                    try:
                        save_self_assess(username, df, st.session_state.get("user_college"))
                        _merge_stu_history(df)
                        st.success(f"✅ 已导入 {len(df)} 条记录并入库")
                    except Exception as e:
                        st.error(f"导入失败：{e}")

    with tabs[2]:
        hist = st.session_state.get("stu_history_df")
        if hist is None or len(hist) == 0:
            hist = load_self_assess(username, college=st.session_state.get("user_college"))
            if hist is not None and len(hist) > 0:
                st.session_state["stu_history_df"] = hist
        if (hist is None or len(hist) == 0) and st.session_state.get("use_demo", False):
            hist = demo_student_history()
            st.session_state["stu_history_df"] = hist
            st.info("当前展示演示自评历史")
        if hist is None or len(hist) == 0:
            st.info("暂无自评数据：请在线自评、导入历史测评表，或打开「演示数据预览」")
            return
        hist = hist.copy()
        dims = [c for c in ["焦虑", "抑郁", "压力", "睡眠障碍", "社交回避"] if c in hist.columns]
        if "自评时间" in hist.columns:
            hist = hist.sort_values("自评时间")
        st.markdown("**📈 各维度趋势**")
        if len(hist) >= 2:
            st.pyplot(chart_multi_trend(hist, dims))
        latest = hist.iloc[-1]
        st.markdown("**🕸️ 最近一次测评雷达图**")
        vals = [float(latest[d] / max(40, hist[d].max()) * 100) for d in dims]
        c1, c2 = st.columns([3, 2])
        with c1:
            st.pyplot(chart_radar(vals, dims, title=f"最近测评（{latest.get('自评时间','')}）"))
        with c2:
            st.markdown("**📋 最近一次得分**")
            st.dataframe(latest.to_frame().T, hide_index=True, use_container_width=True)
            total = float(latest.get("总分", sum(float(latest[d]) for d in dims)))
            if total >= 120:
                lv, col = "⚠️ 高风险", "#D93025"
            elif total >= 80:
                lv, col = "🟡 中风险", "#F57C00"
            else:
                lv, col = "🟢 低风险", "#2E7D32"
            st.markdown(f'<div style="background:{col};color:white;border-radius:12px;padding:14px;font-size:18px;font-weight:bold;text-align:center;">总分 {total:.0f} · {lv}</div>',
                        unsafe_allow_html=True)
            _suggestion_html = {
                "🟢 低风险": ("状态良好", "保持规律作息与适度运动，可每 2-4 周自评一次。", "继续在线自评记录变化"),
                "🟡 中风险": ("建议关注", "近期压力或情绪波动偏大，建议尝试正念减压、与辅导员沟通。", "连续自评 2 次并观察趋势"),
                "⚠️ 高风险": ("建议尽快干预", "请尽快联系学校心理咨询中心或辅导员，安排一对一咨询。", "点击右侧「联系咨询中心」获取帮助"),
            }
            t, d, a = _suggestion_html[lv]
            st.markdown(_suggestion_card_html(t, d, a, color=col), unsafe_allow_html=True)

# ===================== 教师模块 =====================
def run_screen(df_raw, dims, high_thr, mid_thr, batch_name):
    t0 = time.time()
    try:
        with st.spinner("🧠 MCMC 贝叶斯变点采样中（约 5~30 秒，请勿刷新页面）…"):
            df_out, _trace, tau_mean, copula = psycho_risk_screen(
                df_raw, dims_list=dims, high_threshold=high_thr, mid_threshold=mid_thr)
    except Exception as e:
        st.error(f"筛查失败：{e}")
        st.code(traceback.format_exc(), language="python")
        return
    st.session_state["screen_result"] = df_out
    st.session_state["screen_tau"] = tau_mean
    st.session_state["screen_copula"] = copula
    st.session_state["screen_dims"] = dims
    st.session_state["screen_batch_name"] = batch_name
    st.session_state["global_high_thr"] = high_thr
    st.session_state["global_mid_thr"] = mid_thr
    n_h, n_m = int((df_out["风险等级"] == "高危").sum()), int((df_out["风险等级"] == "中危").sum())
    st.success(f"✅ 筛查完成：{len(df_out)} 人，高危 {n_h} 人，中危 {n_m} 人，检测变点 τ≈{tau_mean}，耗时 {time.time()-t0:.1f}s")

def render_teacher_page():
    render_top_bar()
    st.markdown('<div class="sub-title">🧑‍🏫 教师心理普查筛查（贝叶斯变点 + Copula）</div>', unsafe_allow_html=True)
    demo = st.session_state.get("use_demo", False)
    demo_toggle_widget("打开后无需上传即可用模拟普查数据跑通整个筛查流程")
    tabs = st.tabs(["📤 上传与筛查", "📊 结果分析", "📋 学生名单", "📄 导出与入库"])
    with tabs[0]:
        if not _ALG_AVAILABLE:
            st.warning("未检测到 psycho_screen 算法包，请确认 psycho_screen.py 与本文件在同一目录")
        up = st.file_uploader("📤 上传普查量表数据表（含各维度得分列，如：焦虑/抑郁/压力/睡眠障碍/社交回避）",
                              type=["xlsx", "xls", "csv"], key="tea_up")
        df_raw = st.session_state.get("screen_raw_df")
        if up is not None:
            df_raw = parse_upload(up)
            if df_raw is not None:
                st.session_state["screen_raw_df"] = df_raw
        if df_raw is None and demo:
            df_raw = demo_survey()
            st.info("当前使用【演示普查数据】，上传你的表格后自动切换")
        if df_raw is None:
            st.info("请上传普查数据表，或打开「演示数据预览」")
        else:
            st.caption(f"已就绪数据：{len(df_raw)} 行 × {df_raw.shape[1]} 列")
            known = [c for c in KNOWN_DIMS if c in df_raw.columns]
            dims = st.multiselect("🧩 选择量表维度列（用于计算总分与 Copula 相依结构）",
                                  list(df_raw.columns), default=known)
            with st.expander("⚙️ 算法参数（风险阈值）"):
                high_thr = st.slider("高危阈值（综合风险分）", 0.5, 3.0,
                                     float(st.session_state.get("global_high_thr", 1.5)), 0.1, key="tea_high_thr")
                mid_thr = st.slider("中危阈值（综合风险分）", 0.0, 2.5,
                                    float(st.session_state.get("global_mid_thr", 0.5)), 0.1, key="tea_mid_thr")
                batch_name = st.text_input("批次名称", f"普查批次 {datetime.now():%Y-%m-%d}")
            if st.button("🚀 运行智能筛查（MCMC 变点 + Copula）", type="primary",
                         disabled=not (len(dims) >= 1 and _ALG_AVAILABLE), use_container_width=True):
                if not dims:
                    st.warning("请至少选择一个维度列")
                else:
                    run_screen(df_raw, dims, high_thr, mid_thr, batch_name)

    df_res = st.session_state.get("screen_result")
    tau = st.session_state.get("screen_tau")
    dims_res = st.session_state.get("screen_dims", [])
    batch_name = st.session_state.get("screen_batch_name", "")

    with tabs[1]:
        if df_res is None:
            st.info("暂无筛查结果，请先在「上传与筛查」中运行算法")
        else:
            df_res = df_res.copy()
            high_thr = st.slider("高危阈值（调整后即时重分级，无需重跑采样）", 0.5, 3.0,
                                 float(st.session_state.get("global_high_thr", 1.5)), 0.1, key="res_high_thr")
            mid_thr = st.slider("中危阈值（调整后即时重分级）", 0.0, 2.5,
                                float(st.session_state.get("global_mid_thr", 0.5)), 0.1, key="res_mid_thr")
            st.session_state["global_high_thr"] = high_thr
            st.session_state["global_mid_thr"] = mid_thr
            df_res["风险等级"] = classify_from_score(df_res["综合风险分"].values, high_thr, mid_thr)
            n_h, n_m = int((df_res["风险等级"] == "高危").sum()), int((df_res["风险等级"] == "中危").sum())
            c1, c2, c3, c4 = st.columns(4)
            cards = [
                {"title": "筛查人数", "value": f"{len(df_res):,}", "delta": f"变点 τ≈{tau}", "color": "#0F4C99", "icon": "👥"},
                {"title": "高危人数", "value": f"{n_h}", "delta": f"占比 {n_h/len(df_res)*100:.1f}%", "color": "#D93025", "icon": "🚨"},
                {"title": "中危人数", "value": f"{n_m}", "delta": f"占比 {n_m/len(df_res)*100:.1f}%", "color": "#F57C00", "icon": "⚠️"},
                {"title": "低危人数", "value": f"{len(df_res)-n_h-n_m}", "delta": "常规状态", "color": "#2E7D32", "icon": "✅"},
            ]
            for i, c in enumerate([c1, c2, c3, c4]):
                with c:
                    st.markdown(_metric_card_html(**cards[i]), unsafe_allow_html=True)
            st.divider()
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**📊 风险等级分布（饼图）**")
                st.pyplot(chart_risk_pie(df_res))
            with c2:
                st.markdown("**📊 风险等级人数（柱状图）**")
                st.pyplot(chart_risk_bar(df_res))
            st.divider()
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**📍 贝叶斯变点检测（总分随样本变化）**")
                st.pyplot(chart_tau_series(df_res, tau))
            with c2:
                st.markdown("**📈 综合风险分分布**")
                st.pyplot(chart_hist(df_res, "综合风险分"))
            st.divider()
            st.markdown("**🔍 Copula 相依结构（异常分 × 综合风险分）**")
            st.pyplot(chart_copula_scatter(df_res))
            st.divider()
            c1, c2 = st.columns(2)
            with c1:
                if len(dims_res) >= 1:
                    st.markdown("**📦 各维度得分箱线图**")
                    st.pyplot(chart_dims_box(df_res, dims_res))
            with c2:
                if len(dims_res) >= 2:
                    st.markdown("**🔗 维度相关性热力图**")
                    st.pyplot(chart_corr(df_res, dims_res))

    with tabs[2]:
        if df_res is None:
            st.info("暂无筛查结果")
        else:
            levels = st.multiselect("筛选风险等级", ["高危", "中危", "低危"],
                                    default=["高危", "中危"], key="list_levels")
            min_score = st.slider("综合风险分下限", -3.0, 3.0, float(st.session_state.get("global_mid_thr", 0.5)), 0.1, key="list_min")
            view = df_res[df_res["风险等级"].isin(levels) & (df_res["综合风险分"] >= min_score)]
            st.caption(f"共 {len(view)} 人符合条件（全部 {len(df_res)} 人）")
            st.dataframe(view, hide_index=True, use_container_width=True)
            st.download_button("⬇️ 下载筛选名单 CSV", data=view.to_csv(index=False).encode("utf-8-sig"),
                               file_name=f"筛查名单_{batch_name or 'result'}.csv", mime="text/csv")

    with tabs[3]:
        if df_res is None:
            st.info("暂无筛查结果")
        else:
            c1, c2 = st.columns(2)
            with c1:
                if st.button("💾 保存最新批次到数据库", use_container_width=True):
                    try:
                        save_screen_batch(batch_name or f"批次 {datetime.now():%Y-%m-%d}",
                                          st.session_state.get("username", "teacher"),
                                          st.session_state.get("user_college"),
                                          tau if tau is not None else 0,
                                          df_res, st.session_state.get("screen_copula", np.zeros(len(df_res))))
                        st.success("✅ 已保存，管理员/教师下次登录可自动加载")
                    except Exception as e:
                        st.error(f"保存失败：{e}")
            with c2:
                st.caption("导出 Word 报告：含统计表、风险占比图、高危名单")
            stat_df = df_res["风险等级"].value_counts().rename_axis("风险等级").reset_index(name="人数")
            stat_df["占比%"] = (stat_df["人数"] / len(df_res) * 100).round(1)
            buf = generate_word_report(batch_name, tau, df_res, stat_df, chart_risk_pie(df_res))
            if buf:
                st.download_button("📄 下载 Word 筛查报告", data=buf,
                                   file_name=f"筛查报告_{batch_name or 'result'}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            xbuf = io.BytesIO()
            df_res.to_excel(xbuf, index=False)
            xbuf.seek(0)
            st.download_button("⬇️ 下载完整结果 Excel", data=xbuf,
                               file_name=f"筛查结果_{batch_name or 'result'}.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

def generate_word_report(batch_name, tau, df, stat_df, fig_pie):
    if not _DOCX_AVAILABLE:
        st.error("请安装 python-docx：pip install python-docx")
        return None
    doc = Document()
    doc.add_heading("心理普查筛查报告", 0)
    doc.add_heading(f"批次：{batch_name}", level=1)
    doc.add_paragraph(f"检测变点位置：τ ≈ {tau}")
    doc.add_heading("风险统计表", level=2)
    table = doc.add_table(rows=stat_df.shape[0] + 1, cols=stat_df.shape[1])
    for i, col in enumerate(stat_df.columns):
        table.cell(0, i).text = str(col)
    for i, row in stat_df.iterrows():
        for j, val in enumerate(row):
            table.cell(i + 1, j).text = str(val)
    doc.add_heading("风险占比图", level=2)
    img_stream = io.BytesIO()
    fig_pie.savefig(img_stream, format="png")
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(5))
    doc.add_heading("高危学生名单", level=2)
    high_df = df[df["风险等级"] == "高危"].reset_index(drop=True) if "风险等级" in df.columns else pd.DataFrame()
    if len(high_df):
        cols = list(high_df.columns)[:8]
        table2 = doc.add_table(rows=len(high_df) + 1, cols=len(cols))
        for i, col in enumerate(cols):
            table2.cell(0, i).text = str(col)
        for i, row in high_df.iterrows():
            for j, col in enumerate(cols):
                table2.cell(i + 1, j).text = str(row[col])
    else:
        doc.add_paragraph("无高危学生")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ===================== 管理员模块 =====================
def _users_table():
    conn = get_conn()
    df = pd.read_sql("SELECT username, role, college, fullname FROM users", conn)
    conn.close()
    return df

def render_admin_page():
    render_top_bar()
    st.markdown('<div class="sub-title">🛡️ 管理员总控台</div>', unsafe_allow_html=True)
    demo_toggle_widget("打开后可用模拟花名册/台账预览")
    tabs = st.tabs(["📊 总览", "🚨 预警台账", "🧭 干预管理", "👥 用户管理"])

    with tabs[0]:
        up = st.file_uploader("📤 上传学生花名册（学号/姓名/学院/年级/心理状态/最近测评）",
                              type=["xlsx", "xls", "csv"], key="adm_up")
        df_roster = st.session_state.get("admin_roster_df")
        if up is not None:
            df_roster = parse_upload(up)
            if df_roster is not None:
                st.session_state["admin_roster_df"] = df_roster
        if df_roster is None and st.session_state.get("use_demo", False):
            df_roster = demo_roster()
            st.info("当前展示【演示花名册】")
        if df_roster is None:
            st.info("请上传学生花名册，或打开演示数据")
        else:
            n = len(df_roster)
            n_college = df_roster["学院"].nunique() if "学院" in df_roster.columns else 0
            alert_n = int((df_roster["心理状态"] == "预警").sum()) if "心理状态" in df_roster.columns else 0
            focus_n = int((df_roster["心理状态"] == "关注").sum()) if "心理状态" in df_roster.columns else 0
            c1, c2, c3, c4 = st.columns(4)
            cards = [
                {"title": "在校学生总数", "value": f"{n:,}", "delta": f"{n_college} 个学院", "color": "#0F4C99", "icon": "👥"},
                {"title": "预警人数", "value": f"{alert_n}", "delta": f"占比 {alert_n/n*100:.1f}%" if n else "", "color": "#D93025", "icon": "🚨"},
                {"title": "关注人数", "value": f"{focus_n}", "delta": "需重点关注", "color": "#F57C00", "icon": "⚠️"},
                {"title": "状态良好", "value": f"{n-alert_n-focus_n}", "delta": "保持良好", "color": "#2E7D32", "icon": "✅"},
            ]
            for i, c in enumerate([c1, c2, c3, c4]):
                with c:
                    st.markdown(_metric_card_html(**cards[i]), unsafe_allow_html=True)
            st.divider()
            c1, c2 = st.columns(2)
            with c1:
                if "学院" in df_roster.columns:
                    st.markdown("**🏫 学院人数分布**")
                    counts = df_roster["学院"].value_counts()
                    fig, ax = plt.subplots(figsize=(9, 4.5), dpi=110)
                    ax.barh(counts.index[::-1], counts.values[::-1], color="#1a6bc4")
                    for y, v in enumerate(counts.values[::-1]):
                        ax.text(v, y, f" {v}", va="center", fontsize=10)
                    ax.set_xlabel("人数")
                    for sp in ["top", "right"]:
                        ax.spines[sp].set_visible(False)
                    fig.tight_layout()
                    st.pyplot(fig)
            with c2:
                if "年级" in df_roster.columns:
                    st.markdown("**🎓 年级分布**")
                    counts = df_roster["年级"].value_counts()
                    fig, ax = plt.subplots(figsize=(9, 4.5), dpi=110)
                    ax.pie(counts.values, labels=counts.index, autopct="%.1f%%", startangle=90,
                           colors=["#0F4C99", "#1967D2", "#5b9bd5", "#a5c8e8"], textprops={"fontsize": 10})
                    ax.set_title(f"总人数 {n:,}", fontsize=12)
                    fig.tight_layout()
                    st.pyplot(fig)
            if "最近测评" in df_roster.columns:
                st.markdown("**📈 最近测评人数趋势（按月份）**")
                dt = pd.to_datetime(df_roster["最近测评"], errors="coerce").dropna()
                if len(dt):
                    mc = dt.dt.to_period("M").value_counts().sort_index()
                    fig, ax = plt.subplots(figsize=(11, 4), dpi=110)
                    ax.plot([str(m) for m in mc.index], mc.values, marker="o", color="#0F4C99", linewidth=2)
                    ax.set_ylabel("测评人数"); ax.tick_params(axis="x", rotation=30)
                    for sp in ["top", "right"]:
                        ax.spines[sp].set_visible(False)
                    fig.tight_layout()
                    st.pyplot(fig)

    with tabs[1]:
        src = st.radio("数据来源", ["数据库台账", "上传表格", "演示数据"], horizontal=True, key="alert_src")
        df_al = None
        if src == "数据库台账":
            df_al = load_alerts()
            if df_al.empty:
                st.info("数据库暂无预警记录，可在「上传表格」导入或在下方新增")
        elif src == "上传表格":
            up = st.file_uploader("📤 上传预警台账（学号/姓名/学院/风险等级/预警时间/处置状态）",
                                  type=["xlsx", "xls", "csv"], key="alert_up")
            if up is not None:
                df_al = parse_upload(up)
                if df_al is not None:
                    st.session_state["admin_alert_df"] = df_al
            else:
                df_al = st.session_state.get("admin_alert_df")
        else:
            df_al = demo_alert_df()
            st.info("当前展示【演示预警台账】")
        if df_al is not None and not df_al.empty:
            lv = st.multiselect("风险等级", ["高危", "中危", "低危"], default=["高危", "中危"], key="al_lv")
            stt = st.multiselect("处置状态", ["待跟进", "干预中", "已转介", "已结案"], default=[], key="al_stt")
            view = df_al
            if lv and "风险等级" in view.columns:
                view = view[view["风险等级"].isin(lv)]
            if stt and "处置状态" in view.columns:
                view = view[view["处置状态"].isin(stt)]
            st.caption(f"共 {len(view)} 条预警")
            st.dataframe(view, hide_index=True, use_container_width=True)
            st.divider()
            st.markdown("**更新处置状态**")
            c1, c2, c3 = st.columns([3, 2, 1])
            with c1:
                if "id" in view.columns:
                    options = [f"{r['id']}｜{r.get('姓名','')}｜{r.get('风险等级','')}" for _, r in view.iterrows()]
                else:
                    options = [f"{i}｜{r.get('姓名','')}｜{r.get('风险等级','')}" for i, r in view.iterrows()]
                sel = st.selectbox("选择预警记录", options, key="al_sel")
            with c2:
                new_status = st.selectbox("新状态", ["待跟进", "干预中", "已转介", "已结案"], key="al_ns")
            with c3:
                if st.button("更新", use_container_width=True, key="al_update_btn"):
                    idx = options.index(sel)
                    row = view.iloc[idx]
                    if "id" in row:
                        update_alert_status(int(row["id"]), new_status)
                        st.success("✅ 已更新数据库")
                    else:
                        st.session_state["admin_alert_df"].loc[row.name, "处置状态"] = new_status
                        st.success("✅ 已更新（内存表格）")
                    st.rerun()
            with st.expander("➕ 新增预警记录（写入数据库）"):
                with st.form("new_alert"):
                    a1, a2 = st.columns(2)
                    sid = a1.text_input("学号")
                    nm = a2.text_input("姓名")
                    a3, a4 = st.columns(2)
                    col_ = a3.text_input("学院", "计算机学院")
                    lv2 = a4.selectbox("风险等级", ["高危", "中危", "低危"])
                    note = st.text_input("备注（选填）")
                    if st.form_submit_button("保存到数据库"):
                        if sid and nm:
                            save_alert(sid, nm, col_, lv2, note=note)
                            st.success("✅ 已新增")
                            st.rerun()
                        else:
                            st.error("请填写学号和姓名")
        else:
            st.info("暂无预警数据")

    with tabs[2]:
        src = st.radio("数据来源", ["数据库台账", "上传表格", "演示数据"], horizontal=True, key="inter_src")
        df_in = None
        if src == "数据库台账":
            df_in = load_interventions()
            if df_in.empty:
                st.info("数据库暂无干预任务")
        elif src == "上传表格":
            up = st.file_uploader("📤 上传干预台账（学生/学院/风险等级/干预方式/负责人/开始时间/状态）",
                                  type=["xlsx", "xls", "csv"], key="inter_up")
            if up is not None:
                df_in = parse_upload(up)
                if df_in is not None:
                    st.session_state["admin_inter_df"] = df_in
            else:
                df_in = st.session_state.get("admin_inter_df")
        else:
            df_in = demo_intervention_df()
            st.info("当前展示【演示干预台账】")
        if df_in is not None and not df_in.empty:
            stt = st.multiselect("状态", ["待执行", "干预中", "已结案"], default=[], key="in_stt")
            view = df_in if not stt else df_in[df_in["状态"].isin(stt)]
            st.caption(f"共 {len(view)} 条干预任务")
            st.dataframe(view, hide_index=True, use_container_width=True)
            st.divider()
            st.markdown("**更新干预状态**")
            c1, c2, c3 = st.columns([3, 2, 1])
            with c1:
                if "id" in view.columns:
                    options = [f"{r['id']}｜{r.get('学生', '')}｜{r.get('风险等级', '')}" for _, r in view.iterrows()]
                else:
                    options = [f"{i}｜{r.get('学生', '')}｜{r.get('风险等级', '')}" for i, r in view.iterrows()]
                sel = st.selectbox("选择任务", options, key="in_sel")
            with c2:
                new_status = st.selectbox("新状态", ["待执行", "干预中", "已结案"], key="in_ns")
            with c3:
                if st.button("更新", use_container_width=True, key="in_update_btn"):
                    idx = options.index(sel)
                    row = view.iloc[idx]
                    if "id" in row:
                        update_intervention_status(int(row["id"]), new_status)
                        st.success("✅ 已更新数据库")
                    else:
                        st.session_state["admin_inter_df"].loc[row.name, "状态"] = new_status
                        st.success("✅ 已更新（内存表格）")
                    st.rerun()
            with st.expander("➕ 新增干预任务（写入数据库）"):
                with st.form("new_inter"):
                    a1, a2 = st.columns(2)
                    sid = a1.text_input("学号")
                    plan = a2.text_input("干预方式", "一对一咨询")
                    a3, a4 = st.columns(2)
                    col_ = a3.text_input("学院", "计算机学院")
                    handler = a4.text_input("负责人", "张老师")
                    a5, a6 = st.columns(2)
                    start = a5.date_input("开始日期", value=datetime.now().date())
                    end = a6.date_input("结束日期（选填）", value=(datetime.now() + timedelta(days=30)).date())
                    if st.form_submit_button("保存到数据库"):
                        if sid and plan:
                            save_intervention(sid, col_, plan, start.strftime("%Y-%m-%d"),
                                              end.strftime("%Y-%m-%d"), handler)
                            st.success("✅ 已新增")
                            st.rerun()
                        else:
                            st.error("请填写学号和干预方式")
        else:
            st.info("暂无干预数据")

    with tabs[3]:
        st.markdown("**👤 创建新用户**")
        with st.form("new_user"):
            u1, u2 = st.columns(2)
            uname = u1.text_input("用户名")
            pwd = u2.text_input("密码", type="password")
            u3, u4 = st.columns(2)
            role_ = u3.selectbox("角色", ["student", "teacher", "admin"])
            college = u4.text_input("学院（管理员可填全校）", "计算机学院")
            fullname = st.text_input("姓名")
            if st.form_submit_button("创建用户"):
                if uname and pwd:
                    if create_user(uname, pwd, role_, college, fullname):
                        st.success(f"✅ 用户 {uname} 创建成功")
                    else:
                        st.error("用户名已存在")
                else:
                    st.error("用户名和密码不能为空")
        st.markdown("**现有用户**")
        st.dataframe(_users_table(), hide_index=True, use_container_width=True)

# ===================== 主入口 =====================
def main():
    init_session()
    root = st.session_state.get("page_root", "home")
    if root == "home":
        render_home_page()
    elif root == "role_login":
        render_login_page()
    elif root == "app":
        role = st.session_state.get("role", "")
        if role == "student":
            render_student_page()
        elif role == "teacher":
            render_teacher_page()
        elif role == "admin":
            render_admin_page()
        else:
            render_login_page()
    else:
        render_login_page()

if __name__ == "__main__":
    main()
